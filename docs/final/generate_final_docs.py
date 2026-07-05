#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT_DIR = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parent
REPORT_DOCX = OUT_DIR / "INFORME_FINAL_SIMULADOR_SO.docx"
MANUAL_DOCX = OUT_DIR / "MANUAL_INSTALACION_SIMULADOR_SO.docx"
REPORT_MD = OUT_DIR / "informe_final_simulador_so.md"
MANUAL_MD = OUT_DIR / "manual_instalacion_simulador_so.md"
README_MD = OUT_DIR / "README_ENTREGA.md"


MODULE_TABLE = [
    ("src/simulator/simulator_main.c", "Punto de entrada del simulador C; inicializa el estado y procesa comandos."),
    ("src/simulator/simulator.c", "Ciclo de vida global del simulador, creación y limpieza de estructuras."),
    ("src/simulator/input.c", "Interpreta comandos: configuración, procesos manuales, procesos aleatorios, TEST, RUN, PAUSE y STOP."),
    ("src/simulator/process.c", "Construcción de PCBs de usuario, procesos del SO, carga TEST, E/S e interrupciones planificadas."),
    ("src/simulator/process_table.c", "Tabla principal de PCBs; mantiene la propiedad de las estructuras de proceso."),
    ("src/simulator/queue.c", "Colas enlazadas para job, ready, dispositivos, no residentes y terminados."),
    ("src/simulator/scheduler.c", "Planificadores FCFS, SJF, Round Robin y prioridad; variantes expropiativas y no expropiativas."),
    ("src/simulator/dispatcher.c", "Despacho, guardado/restauración de contexto y costo de cambio de contexto."),
    ("src/simulator/cpu.c", "Ejecución por ticks, avance de PC, quantum, interrupciones periódicas, E/S y errores."),
    ("src/simulator/memory.c", "Gestión de memoria física con lista enlazada, bloques de 4 KB, segmentos y estrategias de asignación."),
    ("src/simulator/swap.c", "Planificador de mediano plazo: descarga procesos bloqueados y reingresa procesos listos no residentes."),
    ("src/simulator/io.c", "Colas por dispositivo, bloqueo por E/S y finalización de operaciones."),
    ("src/simulator/interrupt.c", "Registro histórico y contadores de interrupciones por proceso y globales."),
    ("src/simulator/gantt.c", "Construcción del diagrama de Gantt del CPU, inactividad y cambios de contexto."),
    ("src/simulator/protocol.c", "Serialización JSON en líneas SIM_DATA para la interfaz Python."),
    ("src/simulator/names.c", "Nombres legibles de algoritmos, estados, dispositivos, interrupciones y segmentos."),
    ("src/simulator/host_posix.c / host_windows.c", "Funciones dependientes del sistema operativo para tiempo y espera."),
]

PYTHON_TABLE = [
    ("src/gui/main.py", "Entrada de la interfaz PySide6."),
    ("src/gui/windows/*", "Ventanas por algoritmo y ventana principal de selección."),
    ("src/gui/components/*", "Controles visuales: entrada de procesos, mapa de memoria, Gantt, tablas, estadísticas y logs."),
    ("src/gui/services/simulation_service.py", "Servicio de aplicación que coordina comandos, gateway, parser, reducer y estado de sesión."),
    ("src/gui/infrastructure/c_process_gateway.py", "Ejecuta el binario C como subproceso y usa hilos Python para leer stdout/stderr sin bloquear la GUI."),
    ("src/gui/infrastructure/simulator_protocol_parser.py", "Parsea líneas LOG y SIM_DATA emitidas por C."),
    ("src/gui/application/command_serializer.py", "Convierte acciones de la GUI a comandos del simulador C."),
    ("src/gui/state/simulation_state_reducer.py", "Actualiza el estado visible a partir de eventos completos del simulador."),
    ("src/gui/mappers/process_mapper.py", "Mapea PCBs JSON a objetos UiProcess y conserva colores de visualización."),
    ("src/gui/domain/models.py", "Modelos de datos de la interfaz y comandos de proceso."),
]

COMPLIANCE_TABLE = [
    ("Estados NEW, READY, RUNNING, BLOCKED y TERMINATED", "Cumplido", "El PCB usa esos cinco estados visibles; además usa NONE como estado interno antes de la llegada."),
    ("PCB con identificador, estado, PC, SP y datos de planificación", "Cumplido", "El PCB guarda PID, estado, contexto CPU, memoria, datos de scheduler, E/S, interrupciones y errores."),
    ("Procesos del sistema operativo", "Cumplido", "Se crean Kernel, MemoryMgr, Scheduler, IOManager y Swapper con memoria reservada."),
    ("Colas de procesos", "Cumplido", "Existen job_q, ready_q, finished_q, nonresident_q y colas por dispositivo."),
    ("Planificación FCFS", "Cumplido", "Implementado en scheduler.c."),
    ("Planificación SJF", "Cumplido", "Implementado en variantes no expropiativa y expropiativa."),
    ("Planificación Round Robin", "Cumplido", "Implementado con quantum configurable y desalojo por temporizador."),
    ("Planificación por prioridad", "Cumplido", "Implementado en variantes no expropiativa y expropiativa; menor número significa mayor prioridad."),
    ("Dispatcher y cambio de contexto", "Cumplido", "El dispatcher aplica costo configurable, registra Gantt y guarda/restaura contexto."),
    ("Memoria con bloques de 4 KB", "Cumplido", "La memoria se modela como 1 GB físico dividido en frames de 4 KB."),
    ("First Fit, Best Fit y Worst Fit", "Cumplido", "Las tres estrategias seleccionan bloques libres desde la lista enlazada de memoria."),
    ("Segmentos de proceso", "Cumplido", "Se modelan TEXT, DATA, BSS, HEAP y STACK con porcentajes configurables para TEXT, DATA y memoria dinámica."),
    ("Planificador de largo plazo", "Cumplido", "Admite procesos desde job_q cuando llegaron y caben en memoria."),
    ("Planificador de mediano plazo", "Cumplido parcial", "Descarga procesos bloqueados residentes y reingresa READY no residentes; no desaloja READY/RUNNING como víctima."),
    ("Interrupciones", "Cumplido", "Se registran interrupciones de temporizador, syscall, solicitud y fin de E/S, y excepciones simuladas."),
    ("Dispositivos de E/S", "Cumplido", "La simulación usa teclado, disco e impresora para solicitudes aleatorias; la GUI muestra cinco colas de dispositivo."),
    ("Evento de teclado cancelar/continuar", "Cumplido parcial", "El teclado resuelve CONTINUE o CANCEL probabilísticamente; no depende de una tecla real del usuario."),
    ("Errores al 0.5 %", "Cumplido", "Cada proceso puede planificar un error fatal con probabilidad 0.5 %."),
    ("Carga de procesos manuales", "Cumplido", "La GUI y el comando ADD permiten nombre, memoria, burst, llegada, prioridad y segmentos."),
    ("Procesos aleatorios", "Cumplido", "RANDOM permite elegir entre 1 y 20 procesos; se fuerza cobertura mínima de E/S en cargas aleatorias grandes."),
    ("Prueba de 20 procesos", "Cumplido", "El comando TEST genera una carga determinística de 20 procesos."),
    ("Comparación de algoritmos", "Cumplido", "tools/compare_policies.py ejecuta matrices de planificación y memoria y genera un informe HTML."),
    ("Interfaz visual", "Cumplido", "PySide6 muestra Gantt, mapa de memoria, colas, PCB, procesos del SO, estadísticas y log."),
]

COMPARISON_ROWS = [
    ("FCFS", "First Fit", 110.60, 147.91, 70.86, 98.6, 0.0822, 22.63, 0.00, 9),
    ("FCFS", "Best Fit", 110.60, 147.91, 70.86, 98.6, 0.0822, 22.84, 0.00, 9),
    ("FCFS", "Worst Fit", 110.60, 147.91, 70.86, 98.6, 0.0822, 26.50, 0.00, 10),
    ("SJF no expropiativo", "First Fit", 60.13, 99.86, 74.58, 98.5, 0.0859, 54.35, 0.00, 12),
    ("SJF no expropiativo", "Best Fit", 60.41, 100.31, 74.17, 98.5, 0.0859, 51.69, 0.00, 12),
    ("SJF no expropiativo", "Worst Fit", 60.33, 100.31, 74.17, 98.5, 0.0859, 57.14, 0.00, 12),
    ("SJF expropiativo", "First Fit", 56.76, 98.72, 71.78, 98.2, 0.0815, 54.49, 0.00, 12),
    ("SJF expropiativo", "Best Fit", 56.76, 98.72, 71.78, 98.2, 0.0815, 54.23, 0.00, 12),
    ("SJF expropiativo", "Worst Fit", 56.76, 98.72, 71.78, 98.2, 0.0815, 55.22, 0.00, 12),
    ("Round Robin", "First Fit", 101.78, 158.63, 56.40, 97.1, 0.0789, 42.05, 0.00, 13),
    ("Round Robin", "Best Fit", 105.48, 158.92, 51.72, 91.0, 0.0794, 43.85, 0.00, 13),
    ("Round Robin", "Worst Fit", 102.46, 158.66, 55.19, 97.2, 0.0790, 44.73, 0.00, 12),
    ("Prioridad no expropiativa", "First Fit", 61.33, 95.59, 69.39, 98.6, 0.0849, 52.68, 0.00, 11),
    ("Prioridad no expropiativa", "Best Fit", 61.33, 95.59, 69.39, 98.6, 0.0849, 51.02, 0.00, 11),
    ("Prioridad no expropiativa", "Worst Fit", 61.33, 95.59, 69.39, 98.6, 0.0849, 52.31, 0.00, 11),
    ("Prioridad expropiativa", "First Fit", 66.16, 107.07, 71.86, 98.1, 0.0798, 55.18, 0.00, 11),
    ("Prioridad expropiativa", "Best Fit", 66.30, 107.07, 71.86, 98.1, 0.0798, 56.29, 0.00, 11),
    ("Prioridad expropiativa", "Worst Fit", 66.16, 107.07, 71.86, 98.1, 0.0798, 56.00, 0.00, 11),
]


def configure_document(document: Document, landscape: bool = False) -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)


def add_title_page(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Universidad Nacional\nCurso: Sistemas Operativos\n")
    run.bold = True
    run.font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Proyecto: Simulador de Gestión de Procesos y Memoria\nFecha: 05/07/2026")
    document.add_page_break()


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_table(document: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)


def build_report_docx() -> None:
    document = Document()
    configure_document(document, landscape=True)
    add_title_page(document, "Informe Final del Proyecto", "Simulador de Sistemas Operativos")

    document.add_heading("1. Resumen ejecutivo", level=1)
    document.add_paragraph(
        "El proyecto desarrollado es un simulador educativo de gestión de procesos, planificación de CPU, "
        "memoria física, interrupciones y entrada/salida. La simulación principal se ejecuta en C para "
        "mantener el comportamiento del sistema operativo en un núcleo único y determinista. La interfaz "
        "gráfica se implementa en Python con PySide6 y consume los estados completos emitidos por el núcleo C."
    )
    document.add_paragraph(
        "La versión actual permite crear procesos manuales, generar procesos aleatorios, cargar una prueba "
        "determinística de 20 procesos, comparar seis algoritmos de planificación contra tres estrategias "
        "de memoria y visualizar CPU, colas, Gantt, PCB, memoria, procesos del SO, interrupciones y métricas."
    )

    document.add_heading("2. Objetivos", level=1)
    add_bullets(document, [
        "Simular el ciclo de vida de procesos de usuario y procesos del sistema operativo.",
        "Comparar algoritmos de planificación de CPU con métricas de tiempo de espera, retorno, respuesta, throughput y utilización.",
        "Gestionar memoria física con bloques de 4 KB y asignación First Fit, Best Fit y Worst Fit.",
        "Representar PC, SP, segmentos de memoria, colas, interrupciones, E/S y errores dentro del PCB.",
        "Entregar una interfaz visual que muestre el estado del sistema en tiempo de simulación.",
    ])

    document.add_heading("3. Actividades realizadas", level=1)
    add_bullets(document, [
        "Se separó el núcleo de simulación en módulos C especializados por responsabilidad.",
        "Se reemplazó la dependencia de datos inventados en Python por un protocolo donde C emite el estado completo.",
        "Se implementaron los planificadores FCFS, SJF, Round Robin y prioridad, incluyendo variantes expropiativas donde corresponde.",
        "Se implementó una memoria física de 1 GB dividida en frames de 4 KB, con 128 MB reservados para procesos del SO.",
        "Se incorporó la segmentación simulada de cada proceso: TEXT, DATA, BSS, HEAP y STACK.",
        "Se añadió manejo de interrupciones, E/S por dispositivos, errores fatales simulados y registro histórico por PCB.",
        "Se agregó un planificador de mediano plazo para descargar procesos bloqueados y reingresar procesos listos no residentes.",
        "Se creó una GUI PySide6 para configuración, ejecución, visualización y análisis de resultados.",
        "Se prepararon scripts portables de compilación y ejecución para macOS/Linux y Windows.",
        "Se generó una herramienta de comparación automática para evaluar combinaciones de planificación y memoria.",
    ])

    document.add_heading("4. Arquitectura del software", level=1)
    document.add_paragraph(
        "La arquitectura se divide en dos capas principales. La primera capa es el simulador C, que es la autoridad "
        "funcional: crea PCBs, actualiza estados, decide planificación, administra memoria, registra interrupciones "
        "y calcula estadísticas. La segunda capa es la interfaz Python, que ejecuta el binario C como subproceso, "
        "envía comandos por stdin y recibe líneas JSON por stdout."
    )
    document.add_paragraph(
        "El protocolo entre ambas capas usa líneas con prefijo SIM_DATA. Cada mensaje incluye configuración, procesos, "
        "colas, memoria, Gantt y estadísticas. Así la GUI no reconstruye decisiones de simulación; solo serializa "
        "comandos, parsea eventos y representa el estado recibido."
    )
    add_bullets(document, [
        "C: proceso principal de simulación, sin hilos de aplicación.",
        "Python: hilo principal de GUI y dos hilos lectores para stdout/stderr del subproceso C.",
        "Comunicación: stdin/stdout con comandos de texto y JSON por línea.",
        "Persistencia de resultados: HTML comparativo generado por tools/compare_policies.py.",
    ])

    document.add_heading("5. Componentes del núcleo C", level=1)
    add_table(document, ["Módulo", "Responsabilidad"], MODULE_TABLE, [6.0, 19.0])

    document.add_heading("6. Componentes de la interfaz Python", level=1)
    add_table(document, ["Módulo", "Responsabilidad"], PYTHON_TABLE, [7.0, 18.0])

    document.add_heading("7. Modelo de proceso y PCB", level=1)
    document.add_paragraph(
        "Cada proceso se representa con un PCB. El PCB contiene nombre, PID, estado, indicador de proceso del SO, "
        "residencia en memoria, contador de swaps, contexto de CPU, datos de memoria, datos de planificación, "
        "datos de E/S, errores e interrupciones."
    )
    add_bullets(document, [
        "Estados visibles: NEW, READY, RUNNING, BLOCKED y TERMINATED.",
        "Estado interno NONE: usado antes de que el proceso alcance su tiempo de llegada.",
        "Contexto de CPU: program_counter, stack_pointer, pc_offset y sp_offset.",
        "Métricas por proceso: arrival, burst, remaining, start, finish, turnaround, response, ready, blocked, nonresident y cpu_time.",
        "Prioridad: menor número significa mayor prioridad.",
    ])

    document.add_heading("8. Planificación de CPU", level=1)
    document.add_paragraph(
        "El planificador de corto plazo selecciona un proceso de ready_q según el algoritmo activo. El dispatcher "
        "aplica el costo de cambio de contexto configurado, guarda/restaura PC y SP, y registra segmentos de Gantt "
        "para ejecución, inactividad y contexto."
    )
    add_table(document, ["Algoritmo", "Implementación"], [
        ("FCFS", "Extrae el primer proceso de ready_q."),
        ("SJF no expropiativo", "Selecciona el proceso listo con menor tiempo restante y no lo desaloja por otro más corto."),
        ("SJF expropiativo", "Desaloja si aparece un proceso listo con menor remaining_time."),
        ("Round Robin", "Usa FCFS sobre ready_q y desaloja cuando se agota el quantum."),
        ("Prioridad no expropiativa", "Selecciona el menor valor de prioridad y lo mantiene hasta bloqueo o finalización."),
        ("Prioridad expropiativa", "Desaloja si aparece un proceso listo con prioridad numérica menor."),
    ], [6.0, 19.0])

    document.add_heading("9. Gestión de memoria", level=1)
    document.add_paragraph(
        "La memoria se modela como memoria física, no como memoria virtual. El simulador define 1 GB total, frames "
        "de 4 KB y 128 MB reservados para procesos del sistema operativo. La memoria disponible para procesos de "
        "usuario queda representada en una lista enlazada de rangos físicos libres u ocupados."
    )
    document.add_paragraph(
        "Cuando un proceso entra a memoria, se redondea su requerimiento al número de frames necesarios. La lista "
        "enlazada permite buscar un hueco libre con First Fit, Best Fit o Worst Fit. Al liberar memoria se fusionan "
        "huecos contiguos para reducir fragmentación externa."
    )
    add_bullets(document, [
        "TEXT representa el código ejecutable del proceso.",
        "DATA contiene datos inicializados.",
        "BSS representa datos no inicializados y se calcula como parte del área de datos.",
        "HEAP y STACK forman la zona dinámica; crecen progresivamente durante la ejecución simulada.",
        "PC y SP se guardan como offsets antes de descargar memoria y se restauran sobre la nueva base física al reingresar.",
    ])

    document.add_heading("10. Planificación de largo y mediano plazo", level=1)
    document.add_paragraph(
        "El planificador de largo plazo mueve procesos desde job_q a ready_q cuando ya llegaron y existe memoria "
        "suficiente. Si el proceso pide más memoria que la memoria de usuario disponible, termina con error fatal "
        "de tamaño de memoria."
    )
    document.add_paragraph(
        "El planificador de mediano plazo intenta primero reingresar procesos READY no residentes. Si hay procesos "
        "esperando memoria y no existe un hueco suficiente, descarga procesos BLOCKED residentes, priorizando el que "
        "tenga mayor E/S restante y, en empate, mayor tamaño. Esta decisión evita quitar CPU a procesos listos o "
        "ejecutando, pero deja como limitación que no resuelve inanición causada por un proceso residente enorme que "
        "nunca se bloquea."
    )

    document.add_heading("11. Entrada/salida, interrupciones y errores", level=1)
    document.add_paragraph(
        "Los procesos pueden tener una solicitud de E/S configurada. Al alcanzar el punto de E/S, el proceso pasa "
        "a BLOCKED, se registra INT_IO_REQUEST y se encola en el dispositivo correspondiente. Al terminar la E/S "
        "se registra INT_IO_COMPLETE y el proceso vuelve a READY si está residente, o a nonresident_q si terminó "
        "fuera de memoria."
    )
    add_bullets(document, [
        "Dispositivos usados para E/S aleatoria: KEYBOARD, DISK y PRINTER.",
        "La GUI también muestra MOUSE y NETWORK porque el modelo de dispositivos reserva cinco colas.",
        "Interrupciones periódicas: temporizador y llamadas al sistema en función de memoria y burst.",
        "Errores fatales: división entre cero y acceso inválido a memoria con probabilidad de planificación de 0.5 %.",
        "Teclado: al completar E/S puede continuar o cancelar; la cancelación termina el proceso con error fatal.",
    ])

    document.add_heading("12. Configuración del software", level=1)
    add_table(document, ["Parámetro", "Descripción"], [
        ("Planificador", "FCFS, SJF no expropiativo, SJF expropiativo, Round Robin, prioridad no expropiativa o prioridad expropiativa."),
        ("Algoritmo de memoria", "First Fit, Best Fit o Worst Fit."),
        ("Quantum", "Tiempo máximo por turno de CPU para Round Robin."),
        ("Costo de cambio de contexto", "Tiempo agregado por el dispatcher al cambiar de proceso."),
        ("Velocidad", "Factor de velocidad de la simulación visual."),
        ("Proceso manual", "Nombre, CPU burst, memoria, llegada, prioridad y porcentajes de TEXT/DATA/DINÁMICA."),
        ("Procesos aleatorios", "Cantidad entre 1 y 20 generada con parámetros aleatorios."),
        ("TEST", "Carga determinística de 20 procesos para comparar algoritmos."),
        ("Semilla", "SET_SEED permite reproducir escenarios desde consola o herramienta de comparación."),
    ], [6.0, 19.0])

    document.add_heading("13. Comparación de algoritmos", level=1)
    document.add_paragraph(
        "Se ejecutó la herramienta tools/compare_policies.py con el contexto base: switch_cost=0.1, quantum=5, "
        "speed=100, semilla=1234 y carga TEST de 20 procesos. Se evaluaron 18 combinaciones: seis algoritmos de "
        "planificación por tres estrategias de memoria."
    )
    document.add_paragraph(
        "Resultado principal: para esta carga, la mejor combinación de rendimiento fue Prioridad no expropiativa "
        "+ First Fit, con turnaround promedio 95.59 y respuesta promedio 69.39. La mejor combinación de memoria "
        "fue FCFS + First Fit, con fragmentación externa promedio 22.63 % y desperdicio interno 0.00 MB."
    )
    add_table(document, ["Planificación", "Memoria", "Ready", "TAT", "Resp.", "CPU %", "Throughput", "Frag. %", "Waste MB", "Swaps"],
              [(a, b, f"{c:.2f}", f"{d:.2f}", f"{e:.2f}", f"{f:.1f}", f"{g:.4f}", f"{h:.2f}", f"{i:.2f}", j)
               for a, b, c, d, e, f, g, h, i, j in COMPARISON_ROWS],
              [4.8, 2.7, 1.8, 1.8, 1.8, 1.5, 2.0, 1.7, 1.7, 1.4])
    document.add_paragraph(
        "El informe visual generado queda en docs/final/comparisons/comparison_report.html."
    )

    document.add_heading("14. Cumplimiento de requerimientos", level=1)
    add_table(document, ["Requerimiento", "Estado", "Evidencia"], COMPLIANCE_TABLE, [6.0, 3.0, 16.0])

    document.add_heading("15. Librerías y herramientas", level=1)
    add_bullets(document, [
        "Lenguaje C estándar C99 para el núcleo de simulación.",
        "Python 3 para la interfaz y herramientas auxiliares.",
        "PySide6 6.10.3 para la GUI.",
        "python-docx 1.2.0 usado solo para generar este documento final.",
        "pypdf/PyMuPDF usados durante la revisión local de documentos; no son dependencias de ejecución del simulador.",
        "Compiladores soportados: cc/clang/gcc en macOS/Linux; clang/gcc/cl en Windows.",
    ])

    document.add_heading("16. Dificultades funcionales", level=1)
    add_bullets(document, [
        "Mantener a C como fuente única de verdad exigió ampliar el protocolo JSON para enviar memoria, segmentos, interrupciones, E/S, errores y métricas.",
        "La memoria física sin MMU obligó a guardar PC y SP como offsets para poder restaurarlos al reubicar un proceso.",
        "El planificador de mediano plazo debía liberar memoria sin romper la coherencia de colas; se eligió descargar bloqueados como política conservadora.",
        "La GUI debía leer un proceso C en vivo sin congelarse; por eso el gateway usa hilos de lectura en Python.",
        "Las comparaciones requieren escenarios reproducibles; se agregó semilla y una carga TEST fija de 20 procesos.",
    ])

    document.add_heading("17. Limitaciones y mejoras futuras", level=1)
    add_bullets(document, [
        "El planificador de mediano plazo no descarga procesos READY ni RUNNING; sería una mejora para resolver casos de procesos grandes que monopolizan memoria.",
        "La cancelación por teclado se simula con probabilidad, no con una entrada física del usuario durante la ejecución.",
        "El simulador no implementa MMU, paginación ni direcciones virtuales; trabaja con direcciones físicas y offsets.",
        "La fragmentación interna puede ser cero en cargas alineadas a 4 KB; para demostrar desperdicio interno conviene usar tamaños no múltiplos de 4 KB.",
    ])

    document.add_heading("18. Conclusiones", level=1)
    document.add_paragraph(
        "El sistema cumple el objetivo central del proyecto: simular un sistema operativo didáctico con planificación "
        "de CPU, memoria, procesos, interrupciones, E/S, PCB y visualización. La arquitectura separa correctamente "
        "la simulación del renderizado: C decide y Python muestra. Esto permite probar el núcleo desde consola, usar "
        "la GUI para la demostración y ejecutar comparaciones automáticas para sustentar conclusiones."
    )
    document.add_paragraph(
        "Para la exposición se recomienda demostrar primero la compilación con build.sh, luego ejecutar la GUI con "
        "run.sh, cargar TEST o RANDOM 20, alternar algoritmos y mostrar las pestañas de ejecución, PCB, procesos del "
        "SO y estadísticas. Finalmente se puede abrir el informe comparativo HTML para justificar las diferencias "
        "entre políticas."
    )

    document.save(REPORT_DOCX)


def build_manual_docx() -> None:
    document = Document()
    configure_document(document)
    add_title_page(document, "Manual de Instalación y Ejecución", "Simulador de Sistemas Operativos")

    document.add_heading("1. Contenido esperado de la entrega", level=1)
    add_bullets(document, [
        "Código fuente completo del proyecto.",
        "Scripts build.sh y run.sh para macOS/Linux.",
        "Scripts win_build.bat y win_run.bat para Windows.",
        "Archivo requirements.txt con la dependencia PySide6.",
        "Carpeta build con el ejecutable compilado, si se entrega junto al fuente.",
        "Documentación final en docs/final.",
    ])

    document.add_heading("2. Requisitos previos", level=1)
    document.add_heading("2.1 macOS/Linux", level=2)
    add_bullets(document, [
        "Python 3.12 o Python 3 instalado.",
        "Módulo venv de Python disponible.",
        "Compilador C disponible como cc, clang o gcc.",
        "Conexión a internet en la primera ejecución para instalar PySide6 en el entorno virtual local.",
    ])
    document.add_heading("2.2 Windows", level=2)
    add_bullets(document, [
        "Python 3 instalado y disponible como py -3 o python.",
        "Compilador C: clang, gcc o Microsoft cl desde Developer Command Prompt.",
        "Conexión a internet en la primera ejecución para instalar PySide6.",
    ])

    document.add_heading("3. Instalación y compilación en macOS/Linux", level=1)
    add_numbered(document, [
        "Abrir una terminal y entrar a la carpeta del proyecto: cd SimuladorSO",
        "Dar permisos de ejecución si el sistema los perdió al descargar: chmod +x build.sh run.sh",
        "Compilar el núcleo C: ./build.sh",
        "Verificar que aparezca: Compilado: /ruta/al/proyecto/build/simulator",
        "Ejecutar la aplicación gráfica: ./run.sh",
    ])
    document.add_paragraph(
        "El script run.sh crea automáticamente un entorno virtual local en .venv, instala requirements.txt y abre la GUI. "
        "No es necesario activar manualmente un venv externo."
    )

    document.add_heading("4. Instalación y compilación en Windows", level=1)
    add_numbered(document, [
        "Abrir Command Prompt o Developer Command Prompt.",
        "Entrar a la carpeta del proyecto: cd SimuladorSO",
        "Compilar el núcleo C: win_build.bat",
        "Verificar que aparezca build\\simulator.exe.",
        "Ejecutar la aplicación gráfica: win_run.bat",
    ])
    document.add_paragraph(
        "El script win_run.bat crea .venv, instala PySide6 desde requirements.txt y ejecuta gui.main con PYTHONPATH configurado."
    )

    document.add_heading("5. Ejecución desde consola", level=1)
    document.add_paragraph(
        "Además de la GUI, el simulador C puede probarse directamente. Después de compilar, ejecutar build/simulator "
        "en macOS/Linux o build\\simulator.exe en Windows. Los comandos se escriben por stdin."
    )
    add_table(document, ["Comando", "Uso"], [
        ("SET_CONFIG 0 0 0 0.1 100", "Configura FCFS, First Fit, quantum 0, costo 0.1 y velocidad 100."),
        ("SET_CONFIG 3 0 5 0.1 100", "Configura Round Robin, First Fit, quantum 5, costo 0.1 y velocidad 100."),
        ("ADD P1 65536 10 0 3 30 30 40", "Agrega un proceso manual de 64 MB, burst 10, llegada 0 y prioridad 3."),
        ("RANDOM 20", "Carga 20 procesos aleatorios."),
        ("TEST", "Carga la prueba fija de 20 procesos."),
        ("RUN", "Inicia o reanuda la simulación."),
        ("PAUSE", "Pausa la simulación."),
        ("STOP", "Finaliza el proceso simulador."),
    ], [6.0, 10.0])

    document.add_heading("6. Uso básico de la GUI", level=1)
    add_numbered(document, [
        "Ejecutar ./run.sh o win_run.bat.",
        "Elegir el algoritmo de planificación desde la ventana principal.",
        "Seleccionar First Fit, Best Fit o Worst Fit.",
        "Configurar quantum si se usa Round Robin.",
        "Configurar costo de cambio de contexto y velocidad.",
        "Agregar procesos manualmente o usar procesos aleatorios.",
        "Presionar ejecutar y revisar Gantt, mapa de memoria, colas, PCB, procesos del SO y estadísticas.",
    ])

    document.add_heading("7. Generar comparación de algoritmos", level=1)
    document.add_paragraph(
        "La comparación automática recompila el simulador, ejecuta TEST y cruza algoritmos de planificación con "
        "estrategias de memoria. En macOS/Linux:"
    )
    document.add_paragraph(
        ".venv/bin/python tools/compare_policies.py --build --context base --output-dir docs/final/comparisons",
        style=None,
    )
    document.add_paragraph(
        "El resultado se guarda en docs/final/comparisons/comparison_report.html."
    )

    document.add_heading("8. Verificación antes de subir a UNIVIRTUAL", level=1)
    add_numbered(document, [
        "Borrar la carpeta build si se desea probar desde cero.",
        "Ejecutar ./build.sh o win_build.bat.",
        "Confirmar que el ejecutable se crea sin errores.",
        "Ejecutar ./run.sh o win_run.bat.",
        "Cargar TEST y presionar RUN.",
        "Confirmar que el Gantt avanza, que aparecen procesos en PCB y que se actualizan estadísticas.",
        "Ejecutar la comparación base y confirmar que se genera comparison_report.html.",
    ])

    document.add_heading("9. Solución de problemas", level=1)
    add_table(document, ["Problema", "Solución"], [
        ("No se encuentra Python 3", "Instalar Python 3.12 o Python 3 y repetir run.sh/win_run.bat."),
        ("No se puede crear venv en Linux", "Instalar python3-venv y python3-pip. En Debian/Ubuntu: sudo apt install python3-venv python3-pip."),
        ("No se encuentra compilador C", "Instalar Xcode Command Line Tools en macOS, build-essential en Linux o clang/gcc/cl en Windows."),
        ("Qt no inicia en Linux", "Instalar librerías xcb/wayland sugeridas por run.sh o forzar QT_QPA_PLATFORM=xcb/wayland."),
        ("La GUI dice que no existe build/simulator", "Ejecutar primero ./build.sh o win_build.bat."),
        ("No aparecen otros documentos de docs", "El .gitignore mantiene ignorada la documentación local fuera de docs/final para no mezclar borradores con la entrega."),
    ], [5.0, 11.0])

    document.add_heading("10. Archivos principales", level=1)
    add_bullets(document, [
        "src/simulator: núcleo C del simulador.",
        "src/gui: interfaz gráfica PySide6.",
        "tools/compare_policies.py: herramienta de comparación.",
        "build.sh / win_build.bat: compilación del núcleo C.",
        "run.sh / win_run.bat: creación del venv y ejecución de la GUI.",
        "requirements.txt: dependencias Python de ejecución.",
        "docs/final: informe final, manual, fuente Markdown y script generador.",
    ])

    document.save(MANUAL_DOCX)


def render_markdown_table(headers: list[str], rows: Iterable[Iterable]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(str(value).replace("\n", " ") for value in row) + " |")
    return "\n".join(lines)


def build_markdown_sources() -> None:
    report_lines = [
        "# Informe Final del Proyecto - Simulador de Sistemas Operativos",
        "",
        "## Resumen ejecutivo",
        "El proyecto es un simulador educativo de gestión de procesos, planificación de CPU, memoria física, interrupciones y entrada/salida. El núcleo de simulación está implementado en C y la interfaz gráfica en Python con PySide6.",
        "",
        "## Arquitectura",
        "- C concentra la simulación: PCB, colas, CPU, memoria, interrupciones, E/S, Gantt y estadísticas.",
        "- Python ejecuta el simulador C como subproceso y muestra el estado emitido por `SIM_DATA`.",
        "- La GUI usa hilos Python para leer stdout/stderr sin bloquear la ventana.",
        "- La comunicación se realiza con comandos por stdin y JSON por stdout.",
        "",
        "## Componentes C",
        render_markdown_table(["Módulo", "Responsabilidad"], MODULE_TABLE),
        "",
        "## Componentes Python",
        render_markdown_table(["Módulo", "Responsabilidad"], PYTHON_TABLE),
        "",
        "## Cumplimiento de requerimientos",
        render_markdown_table(["Requerimiento", "Estado", "Evidencia"], COMPLIANCE_TABLE),
        "",
        "## Comparación base",
        "Contexto: switch_cost=0.1, quantum=5, speed=100, seed=1234 y carga TEST de 20 procesos.",
        "",
        render_markdown_table(
            ["Planificación", "Memoria", "Ready", "TAT", "Resp.", "CPU %", "Throughput", "Frag. %", "Waste MB", "Swaps"],
            [(a, b, f"{c:.2f}", f"{d:.2f}", f"{e:.2f}", f"{f:.1f}", f"{g:.4f}", f"{h:.2f}", f"{i:.2f}", j)
             for a, b, c, d, e, f, g, h, i, j in COMPARISON_ROWS],
        ),
        "",
        "## Conclusión",
        "La simulación cumple el objetivo central: modelar planificación, memoria, interrupciones, E/S y PCB con visualización. La separación C/Python permite probar el núcleo desde consola, usar la GUI para la demostración y ejecutar comparaciones reproducibles.",
    ]
    REPORT_MD.write_text("\n".join(report_lines) + "\n", encoding="utf-8")

    manual_lines = [
        "# Manual de Instalación - SimuladorSO",
        "",
        "## macOS/Linux",
        "1. `cd SimuladorSO`",
        "2. `chmod +x build.sh run.sh`",
        "3. `./build.sh`",
        "4. `./run.sh`",
        "",
        "`run.sh` crea `.venv`, instala `requirements.txt` y ejecuta la GUI. No se necesita activar un venv externo.",
        "",
        "## Windows",
        "1. `cd SimuladorSO`",
        "2. `win_build.bat`",
        "3. `win_run.bat`",
        "",
        "## Comparaciones",
        "`.venv/bin/python tools/compare_policies.py --build --context base --output-dir docs/final/comparisons`",
        "",
        "## Verificación",
        "- Confirmar que existe `build/simulator` o `build\\simulator.exe`.",
        "- Ejecutar la GUI, cargar `TEST` y presionar `RUN`.",
        "- Revisar Gantt, memoria, PCB, procesos del SO y estadísticas.",
        "- Generar `docs/final/comparisons/comparison_report.html`.",
    ]
    MANUAL_MD.write_text("\n".join(manual_lines) + "\n", encoding="utf-8")

    readme_lines = [
        "# Entrega final",
        "",
        "Archivos generados en esta carpeta:",
        "- `INFORME_FINAL_SIMULADOR_SO.docx`: informe final solicitado.",
        "- `MANUAL_INSTALACION_SIMULADOR_SO.docx`: manual de instalación y ejecución.",
        "- `informe_final_simulador_so.md`: fuente editable del informe.",
        "- `manual_instalacion_simulador_so.md`: fuente editable del manual.",
        "- `generate_final_docs.py`: script reproducible para regenerar los documentos Word.",
        "- `comparisons/comparison_report.html`: comparación de algoritmos de planificación y memoria.",
        "",
        "Nota: `.gitignore` permite versionar `docs/final` y mantiene ignorada la documentación local fuera de esa carpeta.",
    ]
    README_MD.write_text("\n".join(readme_lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_report_docx()
    build_manual_docx()
    build_markdown_sources()
    print(f"Generado: {REPORT_DOCX}")
    print(f"Generado: {MANUAL_DOCX}")
    print(f"Generado: {REPORT_MD}")
    print(f"Generado: {MANUAL_MD}")
    print(f"Generado: {README_MD}")


if __name__ == "__main__":
    main()
